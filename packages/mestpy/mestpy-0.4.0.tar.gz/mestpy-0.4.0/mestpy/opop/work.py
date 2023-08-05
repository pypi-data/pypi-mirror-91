def namesFromExcelToRenameAllFiles(directory, excelFile):
    import os
    import pandas as pd
    import xlsxwriter
    import xlrd
    import openpyxl
    
    files = os.listdir(directory)
    if not files:
        print("В папке нет файлов")
        return
    
    dfIn = pd.read_excel(excelFile, engine='openpyxl')
    dfName=dfIn['Names Files']
    dfNameList = dfName.values.tolist()
    print (len(dfNameList))
    if len(dfNameList) !=  len(files):
            raise TypeError("Количество имен и файлов НЕ равны")
            return    
    

    for ids,f in enumerate(files):
            f_name, f_ext=os.path.splitext(f)
            f_name_new=str(dfNameList[ids])+f_ext
            print("Old Name {}".format(files[ids]))
            print("New Name {}".format(f_name_new))
            os.rename(directory+'/'+f, directory+'/'+f_name_new)

def splitPdfAllPages(pdfFile):
    from PyPDF2 import PdfFileReader, PdfFileWriter, PdfFileMerger
    import os
    f_name, f_ext=os.path.splitext(pdfFile)

    with open(pdfFile, 'rb') as infile:
        reader = PdfFileReader(infile)
        page = reader.numPages
        for i in range(page):
            writerStart = PdfFileWriter()
            writerStart.addPage(reader.getPage(i))
            fpage=f_name+str(i+1)+f_ext
            print(fpage)
            with open(fpage, 'wb') as outfileStart:
                writerStart.write(outfileStart)
                print(f"Страницу файла сохранил как {fpage}")

def splitPdfTwoPart(pdfFile,numBreak):
    import os
    from PyPDF2 import PdfFileReader, PdfFileWriter, PdfFileMerger
    if ((not isinstance(numBreak, int)) or (numBreak<1)):
            raise TypeError("numBreak должно быть целое число не меньше 1")
            return

    with open(pdfFile, 'rb') as infile:
        reader = PdfFileReader(infile)
        page = reader.numPages
        f_name, f_ext=os.path.splitext(pdfFile)

        if (numBreak>page-1):
            raise TypeError("numBreak должно быть целое число не больше страниц файла")
            return
        writerStart = PdfFileWriter()
        writerEnd = PdfFileWriter()
        for i in range(numBreak):
            writerStart.addPage(reader.getPage(i))
        for i in range(numBreak,page):
            writerEnd.addPage(reader.getPage(i))

        print(f"Всего страниц в файле {pdfFile} - {page}")
        print(f"Разбиваю файл после страницы {numBreak}")

        with open('pagesStart.pdf', 'wb') as outfileStart:
            writerStart.write(outfileStart)
        with open('pagesEnd.pdf', 'wb') as outfileEnd:
            writerEnd.write(outfileEnd)
        print(f"Начало файла сохранил как pagesStart.pdf")
        print(f"Конец файла сохранил как pagesEnd.pdf")

def insertStrForFilesName(directory,string,afterNumber):
    import os
    import re
    if ((not isinstance(afterNumber, int)) or (afterNumber<-1)):
            raise TypeError("afterNumber должно быть целое число не меньше -1")
            return

    files = os.listdir(directory)
    if not files:
        print("В папке нет файлов")
        return 
     
    for ids,f in enumerate(files):
        f_name, f_ext=os.path.splitext(f)
        text=f_name
        pattern=""
        repl=str(string)
        if afterNumber==0:
            pattern=r"^"
        elif afterNumber==-1 or afterNumber>len(f):
            pattern=r"$"
        else:
            pat= '.'*int(afterNumber) 
            pattern=re.match(pat,f_name).group(0)
            repl=pattern+str(string)
        m=re.sub(pattern, repl, text)
        print(f"Old Name {text}")
        print(f"New Name {m.strip()}")
        new_name=m.strip()+f_ext
        os.rename(directory+'/'+f, directory+'/'+new_name)

def indexsForFilesWithRE(directory,count,symbols):
    import os
    import re
    if ((not isinstance(count, int)) or (count<1)):
            raise TypeError("count должно быть целое число не меньше 1")
            return
    files = os.listdir(directory)
    if not files:
        print("В папке нет файлов")
        return 
    for ids,f in enumerate(files):
        f_name, f_ext=os.path.splitext(f)
        text=f_name
        pattern=r"^"
        repl=str(ids+1).zfill(count)+symbols
        m=re.sub(pattern, repl, text)
        print(f"Old Name {text}")
        print(f"New Name {m.strip()}")
        new_name=m.strip()+f_ext
        os.rename(directory+'/'+f, directory+'/'+new_name)

def subForFilesWithRE(directory,pattern,repl):
    import os
    import re
    files = os.listdir(directory)
    if not files:
        print("В папке нет файлов")
        return 
    for ids,f in enumerate(files):
        f_name, f_ext=os.path.splitext(f)
        text=f_name
        m=re.sub(pattern, repl, text)
        print(f"Old Name {text}")
        print(f"New Name {m.strip()}")
        new_name=m.strip()+f_ext
        os.rename(directory+'/'+f, directory+'/'+new_name)

def unionPdfsFiles(directory):
    import os
    from PyPDF2 import PdfFileReader, PdfFileMerger
    pdfs = [f for f in os.listdir(directory) if f.endswith(".pdf")]
    if not pdfs:
        print("В папке нет pdf файлов")
        return    
    merger = PdfFileMerger()
    for pdf in pdfs:
        with open(directory+'/'+pdf, 'rb') as fd:
            merger.append(PdfFileReader(fd))
   
    with open('unionPdfsFiles.pdf','wb') as fout:
        merger.write(fout)
    print("Сохранил файл unionPdfsFiles.pdf в основной директории")

def namesFilesToExcel(directoryPdfs):
  from PyPDF2 import PdfFileReader, PdfFileWriter
  import os
  import pandas as pd
  import xlsxwriter
  files = os.listdir(directoryPdfs)
  lNames = [] 
  lExt = []
  for idx, f in enumerate(files):
    f_name, f_ext = os.path.splitext(f)
    lNames.append(f_name)
    lExt.append(f)
  dfToExcel = pd.DataFrame (lNames,columns = ['Names Files'])
  dfToExcel['Full names'] = lExt
  writer = pd.ExcelWriter("namesFiles" + ".xlsx", engine = 'xlsxwriter')
  dfToExcel.to_excel(writer, startrow = 0, startcol = 0, index = False)
  writer.save()
  print("Сохранил файл namesFiles.xlsx в Основной директории")

def docxToPdf(directory):
  from docx2pdf import convert
  import os
  files = [f for f in os.listdir(directory) if f.endswith(".docx")]
  if not files:
    print("В папке нет docx файлов")
    return    
  for ids,f in enumerate(files):
    f_name, f_ext = os.path.splitext(f)
    name = '{} {} {}'.format(ids+1,f_name, f_ext)
    print(name)
  convert(directory+'/')
  print("pdf файлы сохранены в папке c docx")

def check_tables(directory, numTable):
  numTable = numTable-1
  from docx import Document
  import os
  import copy
  files = os.listdir(directory)
  listSortedNatural = files
  document_out = Document() # новый документ с таблицами 0 всех файлов
  text1 = 'таблицы {} из всех docx файлов'.format(numTable)
  document_out.add_paragraph(text1)
  for ids, rpd in enumerate(listSortedNatural): 
    f_name, f_ext = os.path.splitext(rpd)
    if f_ext != ".docx":
      print("Не docx file: {}".format(rpd))
      continue
    text2 = 'Файл '+str(ids+1)
    document_out.add_paragraph(text2)
    document_out.add_paragraph(rpd)  #print(rpd)
    doc_document = os.path.join(directory, rpd)
    document = Document(doc_document)
    for k, table in enumerate(document.tables):
      if k == numTable:
        template0 = document.tables[numTable]
        tbl0 = template0._tbl
        new_tbl0 = copy.deepcopy(tbl0)
        paragraph1 = document_out.add_paragraph()
        paragraph1._p.addnext(new_tbl0)
    document_out.add_page_break()
  document_out.save('all_tables.docx')
  print("создал файл all_tables.docx")

def unionLeftRight(fileForUnion):
  import pandas as pd
  df_rup = pd.read_excel(fileForUnion, usecols = "A,B" )
  dfIn = pd.read_excel(fileForUnion)
  dfu = pd.DataFrame
  dfu = dfIn['Left']+" "+dfIn['Right']
  writer = pd.ExcelWriter("unionLeftRightOut" + ".xlsx", engine = 'xlsxwriter')
  dfu.to_excel(writer)
  writer.save()
  print("Создал Файл unionLeftRightOut.xlsx в корневом каталоге Питон")

def fromRupFullNames(fileRUP):
  import pandas as pd
  df_rup = pd.read_excel(fileRUP, sheet_name = 'Компетенции(2)', index_col = None, header = None )
  df_rup = df_rup.dropna(subset = [4])
  df_rup.drop(df_rup.head(2).index,inplace = True) # drop last n rows 
  df_rup.drop(df_rup.tail(12).index,inplace = True) # drop last n rows 
  df_rup.dropna(axis = 'columns',how = 'all', inplace = True)
  df_rup.drop(df_rup.columns[[0,1, 3,5]], axis = 1, inplace = True)
  df_rup.reset_index(drop = True, inplace = True)
  # df_rup.columns = ['code','id','text']
  df_rup = df_rup.rename(columns = {2: "Index", 4: "Name"})
  writer = pd.ExcelWriter("rupFullNames" + ".xlsx", engine = 'xlsxwriter')
  df_rup.to_excel(writer, sheet_name = 'FullName', startrow = 0, startcol = 0, index = False)
  writer.save()
  print("сохранил файл rupFullNames.xlsx на корневой директории py")

def pagesFromPdfToExcel(directoryPdfs):
  from PyPDF2 import PdfFileReader, PdfFileWriter
  import os
  import pandas as pd
  import xlsxwriter
  pdfs = [f for f in os.listdir(directoryPdfs) if f.endswith(".pdf")]
  if not pdfs:
    print("В папке нет pdf файлов")
    return
  lNames = [] 
  lPages = []
  for idx, f in enumerate(pdfs):
    pdf = PdfFileReader(directoryPdfs+"/"+f)
    page = pdf.numPages
    lPages.append(page)
    f_name, f_ext = os.path.splitext(f)
    lNames.append(f_name)
  dfNamesFiles = pd.DataFrame (lNames,columns = ['Names Files'])
  dfToExcel = pd.DataFrame (lNames,columns = ['Names Files'])
  dfToExcel['Кол-во листов'] = lPages
  writer = pd.ExcelWriter("pagesOfPdfFiles" + ".xlsx", engine = 'xlsxwriter')
  dfToExcel.to_excel(writer, startrow = 0, startcol = 0, index = False)
  writer.save()
  print("Сохранил файл pagesOfPdfFiles.xlsx на корневом каталоге py")


def renameAllFiles(directory,inStrins,outString,numStart,numEnd):
    import os
    files = os.listdir(directory)
    if not files:
        print("В папке нет файлов")

    for ids,f in enumerate(files):
        f_name, f_ext=os.path.splitext(f)
        print("Old Name {}".format(f_name))
        f_name=f_name.replace(inStrins,outString).strip()
        f_start=f_name
        if numStart>0:
            f_start=f_name[:numStart]
        f_end=""
        if numEnd>0:
            f_end=f_name[len(f_name)-numEnd:]
        f_name=f_start+f_end
        print("New Name {}".format(f_name))
        new_name='{}{}'.format(f_name, f_ext)
        os.rename(directory+'/'+f, directory+'/'+new_name)


def indexForFiles(directory,constStr,startNum):
  import os
  files = os.listdir(directory)
  for ids,f in enumerate(files):
    f_name, f_ext = os.path.splitext(f)
    print("Old Name {}".format(f_name))
    new_name = '{}{}'.format(f_name, f_ext)
    new_ids = str(ids+startNum)
    new_name = '{}{} {}{}'.format(constStr,new_ids,f_name, f_ext)
    print("New Name {}".format(new_name))
    os.rename(directory+'/'+f, directory+'/'+new_name)
    
def unionLeftRight(fileForUnion):
  import pandas as pd
  import xlsxwriter
  import xlrd
  import openpyxl
  df_rup = pd.read_excel(fileForUnion, usecols = "A,B", engine='openpyxl' )
  dfIn = pd.read_excel(fileForUnion, engine='openpyxl')
  dfu = pd.DataFrame
  dfu = dfIn['Left']+" "+dfIn['Right']
  writer = pd.ExcelWriter("unionLeftRightOut" + ".xlsx", engine = 'xlsxwriter')
  dfu.to_excel(writer)
  writer.save()
  print("Создал Файл unionLeftRightOut.xlsx в корневом каталоге")
